import os, uuid, tempfile, subprocess
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from pydantic import BaseModel
import markdown2
from bs4 import BeautifulSoup
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from deep_translator import GoogleTranslator
import win32com.client
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from fastapi import APIRouter

router = APIRouter()

@router.get("/word-example")
def sample_word_route():
    return {"message": "Word route is working!"}

app = FastAPI(title="📄 Smart Word AI Pro")
TEMP_DIR = Path(tempfile.gettempdir())

# ✅ Load models locally
print("🔄 Loading AI models locally...")
summarizer = pipeline("summarization", model="google/flan-t5-small", local_files_only=True)
grammar_corrector = pipeline("text2text-generation", model="vennify/t5-base-grammar-correction", local_files_only=True)
classifier = pipeline("text-classification", model="bhadresh-savani/bert-base-uncased-emotion", local_files_only=True)
print("✅ Models loaded.")

def translate_to_english(text: str) -> str:
    try:
        return GoogleTranslator(source='auto', target='en').translate(text)
    except:
        return text

def ensure_word_running():
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = True
        if word.Documents.Count == 0:
            word.Documents.Add()  # नया फाइल खोल दो
        return True
    except Exception as e:
        print("❌ MS Word launch failed:", e)
        return False

def process_text(text: str, summarize=False, grammar_check=False) -> str:
    if summarize and len(text.split()) > 100:
        try:
            text = summarizer(text, max_length=1024, min_length=100, do_sample=False)[0]['summary_text']
        except Exception as e:
            print("Summarizer failed:", e)
    if grammar_check:
        try:
            text = grammar_corrector(text, max_length=1024)[0]['generated_text']
        except Exception as e:
            print("Grammar corrector failed:", e)
    return text

def insert_table(doc: Document, content: str):
    rows = len(content.strip().splitlines())
    table = doc.add_table(rows=rows, cols=2)
    table.style = 'Table Grid'
    for i, line in enumerate(content.strip().splitlines()):
        cells = table.rows[i].cells
        cells[0].text = f"Row {i+1}"
        cells[1].text = line.strip()

def ai_format(doc: Document, content: str):
    lines = content.splitlines()
    for line in lines:
        if line.strip().lower().startswith("heading:"):
            doc.add_heading(line.replace("heading:", "").strip(), level=2)
        elif line.strip().lower().startswith("bold:"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("bold:", "").strip())
            run.bold = True
        elif line.strip().lower().startswith("italic:"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("italic:", "").strip())
            run.italic = True
        else:
            doc.add_paragraph(line.strip())

def detect_type(content: str) -> str:
    result = classifier(content[:512])
    return result[0]['label']

def insert_image(doc: Document):
    try:
        doc.add_picture("sample.jpg", width=Inches(4.0))
    except:
        doc.add_paragraph("[Image missing - 'sample.jpg']")

def generate_doc(title: str, content: str, fmt: str = "text") -> Path:
    doc = Document()
    doc.add_heading(title, level=1)
    doc_type = detect_type(content)
    doc.add_paragraph(f"🧠 Detected as: {doc_type.upper()}")
    if fmt == "markdown":
        html = markdown2.markdown(content)
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup.find_all():
            if tag.name in ["h1", "h2"]:
                doc.add_heading(tag.text.strip(), level=2)
            elif tag.name == "li":
                doc.add_paragraph(tag.text.strip(), style='ListBullet')
            elif tag.name == "p":
                doc.add_paragraph(tag.text.strip())
    else:
        ai_format(doc, content)
        insert_table(doc, content)
        insert_image(doc)

    file_path = TEMP_DIR / f"doc_{uuid.uuid4()}.docx"
    doc.save(file_path)
    return file_path

def get_open_word_content() -> str:
    try:
        word = win32com.client.Dispatch("Word.Application")
        return word.ActiveDocument.Content.Text
    except:
        return ""

def overwrite_open_word_doc(new_content: str):
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.ActiveDocument
        doc.Content.Text = new_content
        doc.Save()
        print("✅ Word document updated in-place.")
    except Exception as e:
        print("❌ Overwrite failed:", e)

from fastapi.responses import HTMLResponse
from fastapi import FastAPI

app = FastAPI()

@app.get("/", response_class=HTMLResponse)
def advanced_ui():
    return HTMLResponse("""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>🧠 Smart Word AI Pro</title>
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        :root {
            --primary-color: #1e88e5;
            --success-color: #4caf50;
            --danger-color: #e53935;
            --border-radius: 8px;
        }
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: #f0f2f5;
            margin: 0;
            padding: 40px;
        }
        .container {
            max-width: 800px;
            margin: auto;
            background: #ffffff;
            border-radius: var(--border-radius);
            box-shadow: 0 8px 20px rgba(0, 0, 0, 0.1);
            padding: 30px 40px;
        }
        h2 {
            color: var(--primary-color);
            margin-bottom: 25px;
            font-weight: 600;
            text-align: center;
        }
        input, textarea, select {
            width: 100%;
            padding: 12px 14px;
            margin-top: 10px;
            border-radius: var(--border-radius);
            border: 1px solid #ccc;
            font-size: 15px;
            box-sizing: border-box;
            transition: border 0.3s ease;
        }
        input:focus, textarea:focus, select:focus {
            border-color: var(--primary-color);
            outline: none;
        }
        .checkbox-group {
            margin: 15px 0;
            display: flex;
            gap: 20px;
        }
        .checkbox-group label {
            font-size: 15px;
        }
        button {
            background-color: var(--primary-color);
            color: #fff;
            padding: 12px 18px;
            font-size: 16px;
            border: none;
            border-radius: var(--border-radius);
            cursor: pointer;
            transition: background-color 0.3s ease;
        }
        button:hover {
            background-color: #1565c0;
        }
        .output {
            margin-top: 25px;
            padding: 16px;
            border-radius: var(--border-radius);
            border-left: 6px solid var(--primary-color);
            background: #f9f9f9;
            font-size: 15px;
        }
        .output.success {
            border-color: var(--success-color);
            color: var(--success-color);
        }
        .output.error {
            border-color: var(--danger-color);
            color: var(--danger-color);
        }
    </style>
</head>
<body>
<div class="container">
    <h2>📄 Smart Word Assistant Pro</h2>
    <input id="title" placeholder="Enter document title..." />
    <textarea id="content" rows="10" placeholder="Paste or write your content here..."></textarea>
    <select id="format">
        <option value="text">Text Format (AI Formatted)</option>
        <option value="markdown">Markdown Format</option>
    </select>
    <div class="checkbox-group">
        <label><input type="checkbox" id="summarize" /> Summarize</label>
        <label><input type="checkbox" id="grammar_check" /> Grammar Check</label>
    </div>
    <button onclick="generateDoc()">📥 Generate Word Document</button>
    <div id="status" class="output" style="display: none;"></div>
</div>
<script>
    async function generateDoc() {
        const statusBox = document.getElementById("status");
        statusBox.style.display = "block";
        statusBox.className = "output";
        statusBox.textContent = "⏳ Processing... Please wait.";

        const formData = new FormData();
        formData.append("title", document.getElementById("title").value);
        formData.append("content", document.getElementById("content").value);
        formData.append("format", document.getElementById("format").value);
        formData.append("summarize", document.getElementById("summarize").checked);
        formData.append("grammar_check", document.getElementById("grammar_check").checked);

        try {
            const res = await fetch('/generate-docx-ui', {
                method: 'POST',
                body: formData
            });

            if (res.ok) {
                const blob = await res.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = "Smart_Document.docx";
                a.click();
                statusBox.textContent = "✅ Document generated and download started.";
                statusBox.classList.add("success");
            } else {
                statusBox.textContent = "❌ Failed to generate document. Please try again.";
                statusBox.classList.add("error");
            }
        } catch (error) {
            statusBox.textContent = "❌ Network error or backend issue.";
            statusBox.classList.add("error");
        }
    }
</script>
</body>
</html>
    """)  # ✅ यह तीन double quotes और एक बंद parenthesis जरूरी है

@app.post("/generate-docx-ui")
def generate_ui(title: str = Form(...), content: str = Form(...), format: str = Form("text"),
                summarize: bool = Form(False), grammar_check: bool = Form(False)):
    processed = process_text(content, summarize, grammar_check)
    file = generate_doc(title, processed, format)
    return FileResponse(file, filename=file.name)

@app.post("/process-open-word/")
def process_live_doc(title: str = Form(...), instruction: str = Form("summarize and fix grammar")):
    content = get_open_word_content()
    if not content:
        return JSONResponse(content={"error": "No MS Word file is open"}, status_code=400)
    task = translate_to_english(instruction)
    summarize = "summarize" in task.lower()
    grammar = "grammar" in task.lower()
    processed = process_text(content, summarize=summarize, grammar_check=grammar)
    overwrite_open_word_doc(processed)
    return JSONResponse(content={"message": "✅ Word document updated in-place."}, status_code=200)

# === CLI / GUI / Server Mode ===
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Smart Word Assistant")
    parser.add_argument("--cli", action="store_true", help="Run in CLI mode")
    parser.add_argument("--gui", action="store_true", help="Run in GUI mode")
    args = parser.parse_args()

    if args.cli:
        print("🧠 Smart Word CLI")
        title = input("📄 Title of document: ")

        mode = input("📥 Input mode - (1) Paste text  (2) Use open Word: ")
        if mode.strip() == "2":
            content = get_open_word_content()
            if not content:
                print("❌ No Word document open.")
                exit()
        else:
            print("✍️ Paste your content below. End with a blank line:")
            lines = []
            while True:
                line = input()
                if line.strip() == "":
                    break
                lines.append(line)
            content = "\n".join(lines)

        tasks = input("⚙️ Tasks to perform (summarize, grammar, both, none): ").lower()
        summarize = "summarize" in tasks
        grammar = "grammar" in tasks

        fmt = input("📑 Format mode (text/markdown): ").strip().lower()
        fmt = fmt if fmt in ["text", "markdown"] else "text"

        overwrite = input("📝 Overwrite open Word file? (y/n): ").strip().lower() == "y"

        processed = process_text(content, summarize=summarize, grammar_check=grammar)

        if overwrite:
            overwrite_open_word_doc(processed)
        else:
            file = generate_doc(title, processed, fmt)
            print(f"✅ Saved: {file}")

    elif args.gui:
        print("🔧 GUI Mode not implemented. Use Web or CLI.")

    else:
        import uvicorn
        uvicorn.run(app, host="0.0.0.0", port=8000)

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--cli", action="store_true")
    args = parser.parse_args()

    ensure_word_running()  # ⬅️ Word को auto-launch करवाएगा

    if args.cli:
        ...
